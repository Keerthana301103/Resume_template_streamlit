import pdfplumber
import docx
from io import BytesIO
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import google.generativeai as genai
import os
from dotenv import load_dotenv




load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
if not api_key: raise ValueError("GEMINI_API_KEY missing.")
genai.configure(api_key=api_key)
model = genai.GenerativeModel('gemini-2.5-flash') 
def extract_text_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text: text += page_text + "\n"
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

# --- Gemini API Interaction ---

def prompt(resume_text):
    template_instruction = """
You are a resume data extractor. Your task is to extract information from the provided resume and curate it as clean, tagged, plain text. 
MUST BE professional throughout and make sure to use Harvard action words, as used in standard resumes, wherever necessary. DO NOT add any special formatting. The Python script will handle all styling

---

FullName: [Full Name]
Designation: [Designation]

ProfessionalOverviewSummary:
[A 2-3 sentence summary of the professional profile, extracted from the resume.Generate based on resume if not explicitly mentioned]

ProfessionalOverviewTable:
Roles | [Summarize Professional  roles  held.Do not repeat same roles.]
Solutions | [Summarize  KEY solution areas, separated by commas. GROUP similar items.]
Industries | [List relevant industries]
Technologies | [Summarize and list KEY 5-7 technologies.GROUP related services.]

KeyEngagementsTable:
Client | Role | Description
[Company Name 1] | [Role at Company1] | [Brief description of engagement 1]

Education:
[Content for the education section]

Publications:
[Content for the publications section]

ProfessionalTrainingCertifications:
[Content for certifications section]

GeographicLocale:
[Content for geographic locale section]


---JOB START---
CompanyName: [Company Name]
Role: [Your Role/Job Title]
Duration: [Start Date – End Date]
Client: [Client Name for the project. If not applicable, write N/A]
Responsibilities:
Make sure the bullet points are concise and follow "Harvard action words" as standard resumes follow.
- [Responsibility point 1]
---JOB END---

Repeat the ---JOB START--- to ---JOB END--- block for each job. If a section is empty, write "None".
"""
    return f"Resume Text:\n{resume_text}\n\n{template_instruction}"

def call_portkey_api(prompt):
    return model.generate_content(prompt).text



def set_table_no_border(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}'); border.set(qn('w:val'), 'nil'); tcBorders.append(border)
            tcPr.append(tcBorders)

def populate_table_cell(cell, heading, content):
    p_heading = cell.paragraphs[0]; run_heading = p_heading.add_run(heading)
    run_heading.font.name = 'Lato'; run_heading.bold = True; run_heading.font.size = Pt(10)
    p_content = cell.add_paragraph(content)
    p_content.style.font.name = 'Lato'; p_content.style.font.size = Pt(9)

def convert_to_docx(text):
    # --- MODIFICATION: Load the template ---
    try:
        doc = Document('template_doc.docx')
    except Exception as e:
        print(f"Error: Could not find or open 'template_doc.docx'. Make sure it's in the same folder.")
        print(f"Details: {e}")
        print("Creating a blank document as a fallback.")
        doc = Document() 

    style = doc.styles['Normal']; font = style.font
    font.name = 'Lato'; font.size = Pt(11)

    resume_data = {}
    lines = text.split('\n'); current_key = None
    
    for line in lines:
        stripped_line = line.strip()

        # Handle JOB markers first, as they don't have colons
        if stripped_line == "---JOB START---":
            current_key = "Jobs"
            if current_key not in resume_data: resume_data[current_key] = []
            resume_data[current_key].append({})
            continue
        elif stripped_line == "---JOB END---":
            current_key = None
            continue

        # Handle lines with colons
        if ":" in line and not stripped_line.startswith('-'):
            key, value = line.split(":", 1)
            key = key.strip()

            if key in ["ProfessionalOverviewSummary", "Education", "Publications", "ProfessionalTrainingCertifications", "GeographicLocale", "ProfessionalOverviewTable", "KeyEngagementsTable"]:
                current_key = key
                resume_data[current_key] = value.strip() + "\n"
            elif current_key == "Jobs" and resume_data.get("Jobs"):
                if key in ["CompanyName", "Role", "Duration", "Client"]:
                    resume_data["Jobs"][-1][key] = value.strip()
                elif key == "Responsibilities":
                    current_key = "Responsibilities"
                    if current_key not in resume_data["Jobs"][-1]:
                        resume_data["Jobs"][-1][current_key] = []
            else: # Simple key-value pairs like FullName
                resume_data[key] = value.strip()
                current_key = None
        
        elif current_key:
            if current_key == "Responsibilities" and resume_data.get("Jobs"):
                resume_data["Jobs"][-1][current_key].append(stripped_line)
            elif current_key in resume_data:
                resume_data[current_key] += line + "\n"
  
    
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_name = p_name.add_run(resume_data.get("FullName", "Candidates Name"))
    run_name.font.color.rgb = RGBColor(204, 31, 32) # Red color
    run_name.font.name = "Lato"
    run_name.bold = True
    run_name.font.size = Pt(18)

    p_des = doc.add_paragraph()
    p_des.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_des = p_des.add_run(resume_data.get("Designation", "Designation(Latest)"))
    run_des.font.color.rgb = RGBColor(0, 0, 0) # Black color
    run_des.font.name = "Lato"
    run_des.bold = False
    run_des.font.size = Pt(12)
    
    doc.add_paragraph()

    doc.add_paragraph("Professional Overview:", style='Heading 2').runs[0].font.color.rgb = RGBColor(204, 31, 32)
    doc.add_paragraph(resume_data.get("ProfessionalOverviewSummary", "").strip())
    
    
    table_lines_po = resume_data.get("ProfessionalOverviewTable", "").strip().split('\n')
    if table_lines_po and '|' in table_lines_po[0]:
       
        table_data = []
        for row_str in table_lines_po:
            if '|' in row_str:
                table_data.append([cell.strip() for cell in row_str.split('|', 1)])

        if table_data:
            table = doc.add_table(rows=len(table_data), cols=2)
            table.style = 'Table Grid'
            
            set_table_no_border(table) 
            table.columns[0].width = Inches(1.5)
            table.columns[1].width = Inches(5.0)
            
            for r, row_data in enumerate(table_data):
                heading, content = row_data
                
                # --- Populate heading cell (col 0) ---
                heading_cell = table.cell(r, 0)
                heading_cell.text = heading
                
                # --- MODIFICATION: Fix spacing & alignment ---
                heading_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                p_heading = heading_cell.paragraphs[0]
                p_heading.runs[0].font.bold = True
                p_heading.runs[0].font.name = 'Lato'
                p_heading.paragraph_format.space_before = Pt(0) # Remove gap
                p_heading.paragraph_format.space_after = Pt(0)  # Remove gap
    
                content_cell = table.cell(r, 1)
                content_cell.text = "" # Clear the default paragraph
                
                # --- MODIFICATION: Fix spacing & alignment ---
                content_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                # --- END MODIFICATION ---

                
                items_list = [item.strip() for item in content.split(',') if item.strip()]
                if not items_list: 
                    items_list = [" "] # Add a space if content is empty

                for i, item in enumerate(items_list):
                    p_bullet = content_cell.paragraphs[0] if i == 0 else content_cell.add_paragraph()
                    
                    # --- MODIFICATION: Fix spacing ---
                    # This ensures all bullet paragraphs (including the first) have no spacing
                    p_bullet.paragraph_format.space_before = Pt(0)
                    p_bullet.paragraph_format.space_after = Pt(0)
                    # --- END MODIFICATION ---

                    run_bullet = p_bullet.add_run('•') # Red bullet
                    run_bullet.font.color.rgb = RGBColor(204, 31, 32) 
                    run_bullet.font.name = 'Lato'
                    
                    p_bullet.add_run('\t') 
                    
                    run_text = p_bullet.add_run(item)
                    run_text.font.name = 'Lato'
                    run_text.font.size = Pt(9)
                    
                    # Hanging indent format
                    p_bullet.paragraph_format.left_indent = Inches(0.25)
                    p_bullet.paragraph_format.first_line_indent = Inches(-0.25)
                    
    doc.add_paragraph()
    

    doc.add_paragraph().add_run("Key Engagements").italic = True
    table_lines_ke = resume_data.get("KeyEngagementsTable", "").strip().split('\n')
    
    if table_lines_ke and '|' in table_lines_ke[0]:
        table_data = [[cell.strip() for cell in row.split('|')] for row in table_lines_ke]
        
        if table_data:
            num_cols = max(len(row) for row in table_data) if table_data else 0
            
            if num_cols > 0:
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Table Grid'
                for r, row_data in enumerate(table_data):
                    for c, cell_data in enumerate(row_data[:num_cols]):
                        table.cell(r, c).text = cell_data

    doc.add_paragraph()

    table_2x2 = doc.add_table(rows=2, cols=2); set_table_no_border(table_2x2)
    populate_table_cell(table_2x2.cell(0, 0), "Education", resume_data.get("Education", "None").strip())
    populate_table_cell(table_2x2.cell(0, 1), "Professional Training/Certifications", resume_data.get("ProfessionalTrainingCertifications", "None").strip())
    populate_table_cell(table_2x2.cell(1, 0), "Publications", resume_data.get("Publications", "None").strip())
    populate_table_cell(table_2x2.cell(1, 1), "Geographic locale", resume_data.get("GeographicLocale", "None").strip())
    doc.add_paragraph()

    doc.add_page_break()
    
    p_exp_heading = doc.add_paragraph("Professional and Business Experience", style='Heading 2')
    p_exp_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
    p_exp_heading.runs[0].font.color.rgb = RGBColor(204, 31, 32)

    for job_data in resume_data.get("Jobs", []):
        p = doc.add_paragraph(); company_run = p.add_run(job_data.get("CompanyName", "")); company_run.font.color.rgb = RGBColor(204, 31, 32); company_run.font.name = 'Lato Black'; company_run.bold = True; company_run.font.size = Pt(12)
        p.add_run('\t'); duration_run = p.add_run(job_data.get("Duration", "")); duration_run.font.name = 'Lato'; p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        p = doc.add_paragraph(); role_run = p.add_run(job_data.get("Role", "")); role_run.font.name = 'Lato'; role_run.bold = True; doc.add_paragraph()
        p = doc.add_paragraph(); client_label_run = p.add_run("CLIENT: "); client_label_run.font.name = 'Lato'; client_label_run.bold = True; client_text_run = p.add_run(job_data.get("Client", "N/A")); client_text_run.font.name = 'Lato'; doc.add_paragraph()
        p = doc.add_paragraph(); resp_run = p.add_run("Responsibilities:"); resp_run.bold=True; resp_run.font.name = 'Lato'; resp_run.underline = True;resp_run.font.size = Pt(11)
        for resp in job_data.get('Responsibilities', []):
            text = resp.lstrip('- ')
    
            p = doc.add_paragraph(style='List Bullet')

            run = p.add_run(text)
    
            run.font.name = 'Lato'
        doc.add_paragraph()
        
    
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer