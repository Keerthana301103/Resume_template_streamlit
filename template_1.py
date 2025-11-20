import pdfplumber
import docx
from io import BytesIO
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from portkey_ai import Portkey
import streamlit as st
import re


M_RED = RGBColor(204, 31, 32)
def clean_pii(text):
    """Removes email addresses and phone numbers using regex."""

    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    
    phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    text = re.sub(email_pattern, '[EMAIL]', text)
    text = re.sub(phone_pattern, '[PHONE]', text)
    
    return text


def extract_text_from_pdf(file):
    """Extracts text from an uploaded PDF file."""
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return clean_pii(text)

def extract_text_from_docx(file):
    """Extracts text from an uploaded DOCX file."""
    doc = docx.Document(file)
    raw_text = "\n".join([para.text for para in doc.paragraphs])
    return clean_pii(raw_text)



def prompt(resume_text):
    """Creates the prompt for the API based on the template."""
    template_instruction = """
You are a resume data extractor. Your task is to extract information from the provided resume and curate it as clean, tagged, plain text. 
MUST BE professional throughout and make sure to use Harvard action words, as used in standard resumes, wherever necessary. DO NOT add any special formatting. The Python script will handle all styling.

---

FullName: [Full Name]

Professional Summary:
[A 2-3 sentence summary of the professional profile, extracted from the resume.Generate based on resume if not explicitly mentioned]

Roles:
[Summarize Professional  roles  held.Do not repeat same roles.]

Technologies:
[Summarize and list KEY 5-7 technologies.GROUP related services. **You MUST preserve the 'Category: Skills' format for each line.** For example: "ETL Tools: Informatica, IICS"]

Education:
[Extract content for the education section.DO NOT Extract percentages/cgpas of the]

Certifications:
[Extract certifications done by the candidate from resume ]

Geographic locale:
[Extract geographic locale from resume section]



---JOB START---
CompanyName: [Company Name, if available. If not, use 'Project']
Role: [Your Role/Job Title]
Duration: [Start Date - End Date]
Client: [Client Name for the project. If not applicable, write N/A]
Description: [Extract the project description]
Responsibilities:
- [Responsibility point 1]
- [Responsibility point 2]

---JOB END---

Repeat the ---JOB START--- to ---JOB END--- block for each job/project. If a section is empty, write "None".
"""
    return f"Resume Text:\n{resume_text}\n\n{template_instruction}"
def call_portkey_api(prompt, portkey_api_key, portkey_base_url):
    """
    Calls the Portkey API with the provided prompt and credentials.
"""
    try:
        portkey = Portkey(
            base_url = portkey_base_url,
            api_key = portkey_api_key
            )
        
        
        response = portkey.chat.completions.create(
            
            model = "@aws-bedrock-use2/us.anthropic.claude-sonnet-4-20250514-v1:0", 
            messages = [
                {"role": "user", "content": prompt}
            ],

)
        
        return response.choices[0].message.content

    except Exception as e:

        st.error(f"Portkey API Error: {str(e)}. Check your Portkey credentials and base_url in Streamlit Secrets.")
        return None

def parse_portkey_text(text):
    """Parses the tagged text from AI into a structured dictionary."""
    resume_data = {"Jobs": []}
    lines = text.split('\n')
    current_key = None
    
    main_keys = ["FullName", "Professional Summary", "Roles", "Technologies", "Education", "Certifications", "Geographic locale", "Professional and Experience Summary"]
    
    for line in lines:
        stripped_line = line.strip()

        if stripped_line == "---JOB START---":
            current_key = "Jobs"
            resume_data["Jobs"].append({})
            continue
        elif stripped_line == "---JOB END---":
            current_key = None
            continue
        
        key_from_line = None
        value_from_line = ""
        
        if ":" in line and not stripped_line.startswith('-'):
            try:
                key_from_line, value_from_line = line.split(":", 1)
                key_from_line = key_from_line.strip()
                value_from_line = value_from_line.strip()
            except ValueError:
                pass 
        
        if key_from_line in main_keys:
            current_key = key_from_line
            resume_data[current_key] = value_from_line
            continue
            
        elif current_key == "Jobs" and key_from_line in ["CompanyName", "Role", "Duration", "Client", "BusinessValue", "Description", "Responsibilities"]:
            if key_from_line in ["CompanyName", "Role", "Duration", "Client", "BusinessValue", "Description"]:
                resume_data["Jobs"][-1][key_from_line] = value_from_line
            elif key_from_line == "Responsibilities":
                current_key = "Responsibilities" 
                if current_key not in resume_data["Jobs"][-1]:
                    resume_data["Jobs"][-1][current_key] = []
                if value_from_line:
                    resume_data["Jobs"][-1][current_key].append(value_from_line)
            continue
            
        elif current_key in main_keys and current_key != "FullName" and stripped_line:
            resume_data[current_key] += "\n" + line.strip()

        elif current_key == "Responsibilities" and stripped_line:
             if "Responsibilities" in resume_data["Jobs"][-1]:
                 resume_data["Jobs"][-1][current_key].append(stripped_line)
                 
    return resume_data

def set_table_no_border(table):
    """Helper function to remove all borders from a table."""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    """Helper to add a styled heading."""
    if not text or text.strip().lower() == 'none':
        return
    style = f'Heading {level}' if level > 0 else 'Title'
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    p.runs[0].font.color.rgb = M_RED
    p.runs[0].font.name = 'Calibri'
    if level == 1:
        p.runs[0].text = text.upper()

def add_content_para(doc, text):
    """Helper to add styled content as a single justified paragraph."""
    if text and text.strip().lower() != 'none':
        para_text = text.strip().replace('\n', ' ')
        if para_text:
            p = doc.add_paragraph(para_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY



def convert_to_docx(text):
    """
    Parses the AI-formatted text and builds the DOCX document.
    """
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    resume_data = parse_portkey_text(text)
    
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = "" 
    
    table_header = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_no_border(table_header)

    cell_left = table_header.cell(0, 0)
    cell_left.width = Inches(1.5)
    p_left = cell_left.paragraphs[0]
    try:
        r_left = p_left.add_run()
        r_left.add_picture('ui\logo.png', width=Inches(1.5))
    except Exception:
        p_left.text = "[logo.png not found]"
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    cell_right = table_header.cell(0, 1)
    cell_right.width = Inches(5.0)
    p_right = cell_right.paragraphs[0]
    run_name = p_right.add_run(resume_data.get("FullName", "Candidate Name"))
    run_name.font.color.rgb = M_RED
    run_name.font.name = 'Calibri'
    run_name.bold = True
    run_name.font.size = Pt(12)
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_heading(doc, "PROFESSIONAL OVERVIEW", level=1)
    
    add_content_para(doc, resume_data.get("Professional Summary"))
    doc.add_paragraph() 

    overview_headings = ["Roles", "Technologies", "Education", "Certifications", "Geographic locale"]
    
    for heading in overview_headings:
        content_text = resume_data.get(heading, "None")
        if content_text.lower() == 'none' or not content_text.strip():
            continue 

        p_heading = doc.add_paragraph()
        run_heading = p_heading.add_run(heading + ":")
        run_heading.font.name = 'Calibri'
        run_heading.bold = True
        run_heading.font.color.rgb = M_RED

        if heading == "Roles" and content_text:
            lines = content_text.strip().split('\n')
            roles = []
            for line in lines:
                roles.extend([r.strip() for r in line.split(',') if r.strip()])
            
            for role in roles:
                doc.add_paragraph(role, style='List Bullet')

        elif heading == "Technologies" and content_text:
            tech_table = doc.add_table(rows=1, cols=2)
            tech_table.style = 'Table Grid' 
            
            tech_table.width = Inches(6.5)
            tech_table.columns[0].width = Inches(2.0)
            tech_table.columns[1].width = Inches(4.5)

            hdr_cells = tech_table.rows[0].cells
            hdr_cells[0].text = '' 
            hdr_cells[1].text = '' 
            hdr_cells[0].paragraphs[0].add_run('Category').bold = True
            hdr_cells[1].paragraphs[0].add_run('Skills').bold = True
            
            for line in content_text.strip().split('\n'):
                if ':' in line:
                    try:
                        category, skills = line.split(':', 1)
                        row_cells = tech_table.add_row().cells
                        row_cells[0].text = category.strip()
                        row_cells[1].text = skills.strip()
                    except ValueError:
                        pass
            doc.add_paragraph() 

        elif heading == "Certifications" and content_text:
            lines = content_text.strip().split('\n')
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line.lstrip('- '), style='List Bullet')

        #
        else:
            for line in content_text.strip().split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph() 

    # --- Professional and Experience Summary Section ---
    doc.add_page_break()
    add_heading(doc, "Professional and Experience Summary", level=1)

    # --- Job/Project Blocks ---
    for i, job_data in enumerate(resume_data.get("Jobs", [])):
        add_heading(doc, f"Project {i+1}", level=2)

        if job_data.get("Client"):
            p = doc.add_paragraph()
            p.add_run("Client: ").bold = True
            p.add_run(job_data.get("Client"))

        if job_data.get("Duration"):
            p = doc.add_paragraph()
            p.add_run("Duration: ").bold = True
            p.add_run(job_data.get("Duration"))

        if job_data.get("Role"):
            p = doc.add_paragraph()
            p.add_run("Role: ").bold = True
            p.add_run(job_data.get("Role"))

        if job_data.get("Description"):
            p = doc.add_paragraph()
            p.add_run("Description: ").bold = True
            add_content_para(doc, job_data.get("Description"))

        responsibilities = job_data.get('Responsibilities', [])
        if responsibilities:
            p = doc.add_paragraph() 
            resp_run = p.add_run("Roles and Responsibilities:")
            resp_run.font.name = 'Calibri'
            resp_run.bold = True
            
            for resp in responsibilities:
                if resp.strip():
                    doc.add_paragraph(resp.lstrip('- '), style='List Bullet')
        
        doc.add_paragraph()
    candidate_name = resume_data.get("FullName", "Candidate_Resume")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, candidate_name